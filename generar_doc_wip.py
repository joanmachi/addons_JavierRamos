from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# --- Estilos globales ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def set_col_width(cell, width_cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_cm * 567)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def shade_cell(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def bold_run(para, text, size=None, color=None):
    run = para.add_run(text)
    run.bold = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run

# --- Título ---
title = doc.add_heading('', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Módulo WIP — De dónde sale cada coste')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x4B, 0x0E, 0x6E)

doc.add_paragraph()

intro = doc.add_paragraph()
intro.add_run(
    'Este documento explica, en lenguaje sencillo, de dónde viene cada número que '
    'aparece en la pantalla de Resumen WIP de Odoo, tanto la columna Teórico como la columna Real, '
    'y dónde hay que ir en Odoo para cambiarlos.'
)

doc.add_paragraph()

# ============================================================
# Sección 1 — MP
# ============================================================
doc.add_heading('1. MP — Materia Prima', level=1)

table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
shade_cell(hdr[0], '4B0E6E')
shade_cell(hdr[1], '4B0E6E')
shade_cell(hdr[2], '4B0E6E')
for i, txt in enumerate(['', 'TEÓRICO', 'REAL']):
    p = hdr[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt)
    r.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r.font.size = Pt(10)

data = [
    ('¿Qué mide?',
     'Lo que deberían costar los materiales si todo va según lo previsto.',
     'Lo que ya se ha gastado o comprometido en materiales hasta hoy.'),
    ('Fuente en Odoo',
     'Lista de materiales (BoM) del producto.',
     'Pedidos de compra recibidos + consumos del operario en taller.'),
    ('Dónde configurarlo',
     'Fabricación → Productos → Listas de materiales',
     'Compras → Pedidos de compra\n(los que tienen la OF en el campo "Fabricación")'),
    ('Cómo se calcula',
     'Cada componente de la BoM × su precio de coste × cantidad a fabricar.',
     'Se toma el MAYOR entre: importe recibido del proveedor O cantidad consumida en taller × precio.'),
]

for i, (concepto, teo, real) in enumerate(data):
    row = table.rows[i + 1].cells
    if i % 2 == 0:
        shade_cell(row[0], 'F3EEF9')
        shade_cell(row[1], 'F3EEF9')
        shade_cell(row[2], 'F3EEF9')
    row[0].paragraphs[0].add_run(concepto).bold = True
    row[1].paragraphs[0].add_run(teo)
    row[2].paragraphs[0].add_run(real)

doc.add_paragraph()
tip = doc.add_paragraph()
tip.add_run('Aviso — BoM incompleta: ').bold = True
tip.add_run(
    'Si el Real supera en más del 50% al Teórico (y el Real > 50 €), Odoo marca la OF con '
    '"BoM incompleta". Significa que en la lista de materiales faltan componentes o '
    'tienen cantidades incorrectas (por ejemplo, servicios externos como mecanizado, '
    'pavonado o pintura que no están en la BoM).'
)

doc.add_paragraph()

# ============================================================
# Sección 2 — MP pendiente de recibir
# ============================================================
doc.add_heading('2. MP pendiente de recibir (solo columna Real)', level=1)

p = doc.add_paragraph()
p.add_run('¿Qué mide? ').bold = True
p.add_run('Material que ya está pedido al proveedor pero que todavía no ha llegado al almacén.')

p = doc.add_paragraph()
p.add_run('Fuente en Odoo: ').bold = True
p.add_run('Pedidos de compra en estado Confirmado, con la OF marcada en el campo "Fabricación", '
          'donde la cantidad recibida es menor que la pedida.')

p = doc.add_paragraph()
p.add_run('Dónde verlo: ').bold = True
p.add_run('Compras → Pedidos de compra → buscar por referencia de la OF en el campo Fabricación.')

doc.add_paragraph()

# ============================================================
# Sección 3 — Tiempo de fabricación
# ============================================================
doc.add_heading('3. Tiempo de fabricación', level=1)

table2 = doc.add_table(rows=4, cols=3)
table2.style = 'Table Grid'
hdr2 = table2.rows[0].cells
for c in hdr2:
    shade_cell(c, '1F497D')
for i, txt in enumerate(['', 'TEÓRICO', 'REAL']):
    p2 = hdr2[i].paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(txt)
    r2.bold = True
    r2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r2.font.size = Pt(10)

data2 = [
    ('¿Qué mide?',
     'Minutos previstos en las operaciones del producto.',
     'Minutos reales fichados por los operarios en taller.'),
    ('Fuente en Odoo',
     'Operaciones del producto (routing).',
     'Fichajes de entrada/salida del operario en la vista Taller.'),
    ('Dónde configurarlo',
     'Fabricación → Configuración → Operaciones → campo "Duración prevista (min)"',
     'Se registra automáticamente cuando el operario ficha en Taller.'),
]

for i, (concepto, teo, real) in enumerate(data2):
    row = table2.rows[i + 1].cells
    if i % 2 == 0:
        shade_cell(row[0], 'EEF3FB')
        shade_cell(row[1], 'EEF3FB')
        shade_cell(row[2], 'EEF3FB')
    row[0].paragraphs[0].add_run(concepto).bold = True
    row[1].paragraphs[0].add_run(teo)
    row[2].paragraphs[0].add_run(real)

doc.add_paragraph()

# ============================================================
# Sección 4 — Coste operario
# ============================================================
doc.add_heading('4. Coste operario', level=1)

table3 = doc.add_table(rows=4, cols=3)
table3.style = 'Table Grid'
hdr3 = table3.rows[0].cells
for c in hdr3:
    shade_cell(c, '375623')
for i, txt in enumerate(['', 'TEÓRICO', 'REAL']):
    p3 = hdr3[i].paragraphs[0]
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(txt)
    r3.bold = True
    r3.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r3.font.size = Pt(10)

data3 = [
    ('¿Qué mide?',
     'Lo que costaría pagar al operario si tarda lo previsto.',
     'Lo que ha costado pagar al operario por el tiempo realmente fichado.'),
    ('Cómo se calcula',
     'Minutos previstos ÷ 60 × tarifa €/hora del operario.',
     'Minutos fichados ÷ 60 × tarifa €/hora del operario.'),
    ('Dónde configurar la tarifa',
     '1.º Ficha del empleado: RRHH → Empleados → campo "Coste/hora"\n'
     '2.º Si está a 0, usa el centro de trabajo: Fabricación → Configuración → Centros de trabajo → "Coste/hora del empleado"',
     'Igual que el teórico.'),
]

for i, (concepto, teo, real) in enumerate(data3):
    row = table3.rows[i + 1].cells
    if i % 2 == 0:
        shade_cell(row[0], 'EFF5EC')
        shade_cell(row[1], 'EFF5EC')
        shade_cell(row[2], 'EFF5EC')
    row[0].paragraphs[0].add_run(concepto).bold = True
    row[1].paragraphs[0].add_run(teo)
    row[2].paragraphs[0].add_run(real)

p_tip = doc.add_paragraph()
p_tip.add_run('Por qué sale 0 € en pantalla: ').bold = True
p_tip.add_run(
    'Los campos "Coste/hora del empleado" en los centros de trabajo y "Coste/hora" en las '
    'fichas de empleados están a 0. Para que aparezca un coste hay que rellenarlos.'
)

doc.add_paragraph()

# ============================================================
# Sección 5 — Coste máquina
# ============================================================
doc.add_heading('5. Coste máquina', level=1)

table4 = doc.add_table(rows=4, cols=3)
table4.style = 'Table Grid'
hdr4 = table4.rows[0].cells
for c in hdr4:
    shade_cell(c, '7F3F00')
for i, txt in enumerate(['', 'TEÓRICO', 'REAL']):
    p4 = hdr4[i].paragraphs[0]
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run(txt)
    r4.bold = True
    r4.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r4.font.size = Pt(10)

data4 = [
    ('¿Qué mide?',
     'Lo que costaría el uso de la máquina/centro de trabajo si se tarda lo previsto.',
     'Lo que ha costado el uso de la máquina por el tiempo realmente fichado.'),
    ('Cómo se calcula',
     'Minutos previstos ÷ 60 × coste/hora de la máquina.',
     'Minutos fichados ÷ 60 × coste/hora de la máquina.'),
    ('Dónde configurar la tarifa',
     'Fabricación → Configuración → Centros de trabajo → campo "Coste por hora"',
     'Igual que el teórico.'),
]

for i, (concepto, teo, real) in enumerate(data4):
    row = table4.rows[i + 1].cells
    if i % 2 == 0:
        shade_cell(row[0], 'FDF3E3')
        shade_cell(row[1], 'FDF3E3')
        shade_cell(row[2], 'FDF3E3')
    row[0].paragraphs[0].add_run(concepto).bold = True
    row[1].paragraphs[0].add_run(teo)
    row[2].paragraphs[0].add_run(real)

doc.add_paragraph()

# ============================================================
# Sección 6 — TOTAL
# ============================================================
doc.add_heading('6. TOTAL', level=1)

p = doc.add_paragraph()
p.add_run('Es la suma de los tres costes:').bold = True

formula = doc.add_paragraph(style='Normal')
formula.paragraph_format.left_indent = Cm(1)
r = formula.add_run('TOTAL = MP  +  Coste operario  +  Coste máquina')
r.font.name = 'Courier New'
r.font.size = Pt(11)
r.bold = True

doc.add_paragraph()
doc.add_paragraph(
    'La columna Teórico es lo que se espera gastar si se siguen exactamente la lista de materiales '
    'y los tiempos del routing. La columna Real es lo que ya se ha gastado o comprometido hasta el momento.'
)

doc.add_paragraph()

# ============================================================
# Sección 7 — Resumen de configuración
# ============================================================
doc.add_heading('7. Resumen: dónde configurar cada cosa en Odoo', level=1)

table5 = doc.add_table(rows=6, cols=2)
table5.style = 'Table Grid'
hdr5 = table5.rows[0].cells
shade_cell(hdr5[0], '2E2E2E')
shade_cell(hdr5[1], '2E2E2E')
for i, txt in enumerate(['Si quiero cambiar...', 'Voy a...']):
    p5 = hdr5[i].paragraphs[0]
    r5 = p5.add_run(txt)
    r5.bold = True
    r5.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

config_data = [
    ('Lo que cuesta un componente (materia prima)',
     'Inventario o Compras → Productos → pestaña "Información general" → Precio de coste'),
    ('Qué componentes entran en el producto',
     'Fabricación → Productos → Listas de materiales'),
    ('Cuánto tiempo se prevé para fabricar',
     'Fabricación → Configuración → Operaciones → Duración prevista (min)'),
    ('Cuánto cobra la máquina por hora',
     'Fabricación → Configuración → Centros de trabajo → Coste por hora'),
    ('Cuánto cobra el operario por hora',
     'Fabricación → Configuración → Centros de trabajo → Coste/hora del empleado\nO bien: RRHH → Empleados → Coste/hora'),
]

for i, (q, a) in enumerate(config_data):
    row = table5.rows[i + 1].cells
    if i % 2 == 0:
        shade_cell(row[0], 'F5F5F5')
        shade_cell(row[1], 'F5F5F5')
    row[0].paragraphs[0].add_run(q).bold = True
    row[1].paragraphs[0].add_run(a)

# ============================================================
# Pie
# ============================================================
doc.add_paragraph()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer_p.add_run('Documento generado automáticamente — Módulo apunts_jr_wip_costes_of — Javier Ramos')
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

output = r"C:\Users\Andrea Sevilla\Documents\WIP_Costes_Explicacion.docx"
doc.save(output)
print(f"Guardado en: {output}")
